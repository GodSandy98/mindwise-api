import io
import json
import threading
import zipfile
from datetime import datetime
from urllib.parse import quote
from typing import List, Dict
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel
from openai import OpenAI
from sqlalchemy.orm import Session
from sqlalchemy import func

from app.core.config import QWEN_API_KEY, QWEN_BASE_URL, QWEN_MODEL
from app.db.session import get_db
from app.models.indicator import Indicator
from app.models.indicator_text import IndicatorText
from app.models.score_student import ScoreStudent
from app.models.report import Report
from app.models.report_indicator import ReportIndicator
from app.models.persona_template import PersonaTemplate
from app.models.student import Student
from app.models.class_ import Class
from app.models.teacher import Teacher
from app.models.batch_job import BatchJob
from app.models.exam import Exam
from app.schemas.report import (
    ReportGenerateRequest,
    ReportGenerateResponse,
    IndicatorAnalysis,
    SystemLevelResult,
    PersonaResult,
    ReportSaveRequest,
    ReportSaveResponse,
    ReportGetResponse,
    SavedIndicatorAnalysis,
    IndicatorHistoryResponse,
    IndicatorHistory,
    IndicatorVersion,
    BatchGenerateRequest,
    StudentReportStatus,
)
from app.api.v1.deps import get_current_teacher, require_psych_or_above, assert_student_class_access

router = APIRouter(prefix="/reports", tags=["reports"])

_LEVEL_THRESHOLDS = {"H": 0.67, "L": -0.67}
_LEVEL_LABELS = {"H": "优势", "M": "中等", "L": "待提升"}


def _score_to_level(score: float) -> str:
    if score >= _LEVEL_THRESHOLDS["H"]:
        return "H"
    if score <= _LEVEL_THRESHOLDS["L"]:
        return "L"
    return "M"


def _build_llm_prompt(
    top3: List[dict],
    bottom3: List[dict],
    top3_templates: Dict[int, dict],
    bottom3_templates: Dict[int, dict],
) -> str:
    """Build LLM prompt with template examples for personalized generation."""

    def _fmt_indicator(item: dict, template: dict | None, is_strength: bool) -> str:
        lines = [f"- {item['name']}（z分数: {item['score']:.2f}, 等级: {_LEVEL_LABELS[item['level']]}）"]
        if template:
            lines.append(f"  参考话术样例（学生版）：{template.get('analysis_student', '')}")
            lines.append(f"  参考话术样例（教师版）：{template.get('analysis_teacher', '')}")
            if not is_strength and template.get('suggestion_student'):
                lines.append(f"  参考建议样例：{template.get('suggestion_student', '')}")
        return "\n".join(lines)

    strengths_text = "\n".join(
        _fmt_indicator(item, top3_templates.get(item["indicator_id"]), True)
        for item in top3
    )
    weaknesses_text = "\n".join(
        _fmt_indicator(item, bottom3_templates.get(item["indicator_id"]), False)
        for item in bottom3
    )

    return f"""你是一位专业的学生心理测评报告撰写专家。请根据以下学生的心理测评数据，生成个性化的分析报告。

【注意】下方每个指标附带了"参考话术样例"，它们是该指标在该等级下的通用模板。请以此为基础，结合该学生的具体得分特征，生成更加个性化、有针对性的内容。不要照搬模板，要体现这个学生的独特性。

【优势指标（得分最高的三项）】
{strengths_text}

【不足指标（得分最低的三项）】
{weaknesses_text}

请严格按照以下 JSON 格式输出，不要输出任何 JSON 以外的内容：

{{
  "summary": "<综合概述，150-200字，用第二人称'你'，学生口吻，温暖鼓励。先概括最突出的优势（结合优势指标第一项），再点出最需要关注的成长空间（结合不足指标第一项），最后给出一句整体鼓励。不要列举所有指标，只抓最核心的一优一劣做自然过渡的叙述>",
  "strengths": [
    {{
      "indicator_name": "<指标名称>",
      "analysis": "<学生版分析，100-150字，第二人称'你'，温暖共情，基于样例但个性化>",
      "analysis_teacher": "<教师版分析，50-100字，专业直接，面向教师的诊断性描述>"
    }}
  ],
  "weaknesses": [
    {{
      "indicator_name": "<指标名称>",
      "analysis": "<学生版分析，100-150字，第二人称'你'，客观有建设性，基于样例但个性化>",
      "analysis_teacher": "<教师版分析，50-100字，专业直接，面向教师的诊断性描述>",
      "suggestion": "<学生版改进建议，50-80字，具体可操作>"
    }}
  ]
}}

要求：
1. summary 是整篇报告的开篇综合概述，学生口吻，温暖鼓励，先扬后抑，自然流畅
2. strengths 必须包含 3 条，对应三个优势指标，顺序与输入一致
3. weaknesses 必须包含 3 条，对应三个不足指标，顺序与输入一致
4. 学生版语言风格：温暖、共情、去标签化、去病理化，用"你"
5. 教师版语言风格：专业、直接、诊断性，便于教师快速了解学生状况
6. 建议要具体可操作，不要空泛的"加油"类话语
7. 不要照搬参考样例，要根据学生的具体z分数和等级生成个性化内容
8. 只输出合法 JSON，不要有任何多余文字"""


def _generate_for_student(db: Session, student_id: int, exam_id: int) -> ReportGenerateResponse:
    """Generate report for a single student using template-guided LLM."""
    if not QWEN_API_KEY:
        raise HTTPException(status_code=500, detail="QWEN_API_KEY 未配置")

    # 1. Fetch leaf indicator scores
    leaf_indicators = db.query(Indicator).filter(Indicator.is_leaf == 1).all()
    leaf_ids = {ind.id for ind in leaf_indicators}
    indicators_by_id = {ind.id: ind for ind in leaf_indicators}

    rows = (
        db.query(ScoreStudent)
        .filter(
            ScoreStudent.student_id == student_id,
            ScoreStudent.exam_id == exam_id,
            ScoreStudent.indicator_id.in_(leaf_ids),
        )
        .all()
    )

    if not rows:
        raise HTTPException(
            status_code=404,
            detail=f"未找到 student_id={student_id}, exam_id={exam_id} 的得分数据，请先调用 /scores/compute",
        )

    # 2. Compute system levels
    system_zscores: Dict[str, List[float]] = defaultdict(list)
    for r in rows:
        ind = indicators_by_id.get(r.indicator_id)
        if ind and ind.system and r.score_standardized is not None:
            system_zscores[ind.system].append(r.score_standardized)

    system_levels = {}
    for system, scores in system_zscores.items():
        avg_z = sum(scores) / len(scores) if scores else 0.0
        system_levels[system] = (round(avg_z, 4), _score_to_level(avg_z))

    mot_level = system_levels.get("motivation", (0, "M"))[1]
    reg_level = system_levels.get("regulation", (0, "M"))[1]
    exe_level = system_levels.get("execution", (0, "M"))[1]
    persona_code = f"{mot_level}{reg_level}{exe_level}"

    # 3. Match persona template (this stays template-based, no LLM needed)
    persona_tmpl = db.query(PersonaTemplate).filter(PersonaTemplate.code == persona_code).first()
    if not persona_tmpl:
        raise HTTPException(status_code=500, detail=f"未找到画像模板: {persona_code}")

    persona = PersonaResult(
        code=persona_code,
        teacher_label=persona_tmpl.teacher_label,
        teacher_description=persona_tmpl.teacher_description,
        student_label=persona_tmpl.student_label,
        student_description=persona_tmpl.student_description,
    )

    # 4. Sort and pick top3 / bottom3
    scored = sorted(
        [
            {
                "indicator_id": r.indicator_id,
                "name": indicators_by_id[r.indicator_id].name,
                "system": indicators_by_id[r.indicator_id].system or "",
                "score_raw": r.score_raw,
                "score": r.score_standardized,
                "level": _score_to_level(r.score_standardized),
            }
            for r in rows
            if r.score_standardized is not None and r.indicator_id in indicators_by_id
        ],
        key=lambda x: x["score"],
        reverse=True,
    )

    if len(scored) < 6:
        raise HTTPException(status_code=422, detail="有效指标得分不足6项，无法生成报告")

    top3 = scored[:3]
    bottom3 = scored[-3:][::-1]  # lowest 3, from low to high

    # 5. Load template texts as examples for LLM
    all_ids = [item["indicator_id"] for item in top3 + bottom3]
    all_levels = {item["indicator_id"]: item["level"] for item in top3 + bottom3}
    indicator_texts = db.query(IndicatorText).filter(
        IndicatorText.indicator_id.in_(all_ids)
    ).all()

    def _build_template_lookup(items: List[dict]) -> Dict[int, dict]:
        lookup = {}
        for item in items:
            ind_id = item["indicator_id"]
            level = item["level"]
            student_t = next((t for t in indicator_texts if t.indicator_id == ind_id and t.level == level and t.view == "student"), None)
            teacher_t = next((t for t in indicator_texts if t.indicator_id == ind_id and t.level == level and t.view == "teacher"), None)
            lookup[ind_id] = {
                "analysis_student": student_t.analysis if student_t else "",
                "analysis_teacher": teacher_t.analysis if teacher_t else "",
                "suggestion_student": student_t.suggestion if student_t else "",
                "suggestion_teacher": teacher_t.suggestion if teacher_t else "",
                "golden_quote": student_t.golden_quote if student_t else "",
            }
        return lookup

    top3_templates = _build_template_lookup(top3)
    bottom3_templates = _build_template_lookup(bottom3)

    # 6. Call LLM with template-guided prompt
    prompt = _build_llm_prompt(top3, bottom3, top3_templates, bottom3_templates)

    try:
        client = OpenAI(api_key=QWEN_API_KEY, base_url=QWEN_BASE_URL)
        response = client.chat.completions.create(
            model=QWEN_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "你是专业的学生心理测评报告撰写专家，擅长将测评数据转化为温暖、个性化的文字分析。你需要参考给定的话术模板，但生成更有针对性的个性化内容。",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.7,
            response_format={"type": "json_object"},
        )
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"调用 LLM API 失败: {str(e)}")

    # 7. Parse response
    try:
        content = response.choices[0].message.content
        data = json.loads(content)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"解析 LLM 响应失败: {str(e)}")

    # 8. Assemble results
    top3_map = {item["name"]: item for item in top3}
    bottom3_map = {item["name"]: item for item in bottom3}

    strengths = []
    for item in data.get("strengths", []):
        name = item["indicator_name"]
        meta = top3_map.get(name)
        if not meta:
            # Fallback: match by position
            idx = len(strengths)
            meta = top3[idx] if idx < len(top3) else top3[0]
        strengths.append(IndicatorAnalysis(
            indicator_id=meta["indicator_id"],
            indicator_name=name,
            score_raw=meta["score_raw"],
            score_standardized=meta["score"],
            level=meta["level"],
            system=meta["system"],
            analysis=item.get("analysis", ""),
            analysis_teacher=item.get("analysis_teacher", ""),
            suggestion=None,
        ))

    weaknesses = []
    for item in data.get("weaknesses", []):
        name = item["indicator_name"]
        meta = bottom3_map.get(name)
        if not meta:
            idx = len(weaknesses)
            meta = bottom3[idx] if idx < len(bottom3) else bottom3[0]
        weaknesses.append(IndicatorAnalysis(
            indicator_id=meta["indicator_id"],
            indicator_name=name,
            score_raw=meta["score_raw"],
            score_standardized=meta["score"],
            level=meta["level"],
            system=meta["system"],
            analysis=item.get("analysis", ""),
            analysis_teacher=item.get("analysis_teacher", ""),
            suggestion=item.get("suggestion"),
        ))

    system_level_results = [
        SystemLevelResult(system=sys, avg_z=round(vals[0], 4), level=vals[1])
        for sys, vals in system_levels.items()
    ]

    summary = data.get("summary", "")

    return ReportGenerateResponse(
        student_id=student_id,
        exam_id=exam_id,
        persona=persona,
        system_levels=system_level_results,
        summary=summary,
        strengths=strengths,
        weaknesses=weaknesses,
    )


@router.post("/generate", response_model=ReportGenerateResponse)
def generate_report(
    payload: ReportGenerateRequest,
    db: Session = Depends(get_db),
    current: Teacher = Depends(require_psych_or_above),
) -> ReportGenerateResponse:
    """Generate personalized report: persona (template) + top3/bottom3 (LLM with template examples)."""
    student = db.query(Student).filter(Student.id == payload.student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    assert_student_class_access(current, student.class_id)

    return _generate_for_student(db, payload.student_id, payload.exam_id)


@router.post("/save", response_model=ReportSaveResponse)
def save_report(
    payload: ReportSaveRequest,
    db: Session = Depends(get_db),
    current: Teacher = Depends(get_current_teacher),
):
    student = db.query(Student).filter(Student.id == payload.student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    assert_student_class_access(current, student.class_id)

    # Resolve persona template
    persona_tmpl = db.query(PersonaTemplate).filter(
        PersonaTemplate.code == payload.persona_code
    ).first()

    # Build name -> id mapping
    all_names = [i.indicator_name for i in payload.strengths] + [i.indicator_name for i in payload.weaknesses]
    indicators = db.query(Indicator).filter(Indicator.name.in_(all_names)).all()
    name_to_id = {ind.name: ind.id for ind in indicators}

    try:
        with db.begin_nested():
            report = (
                db.query(Report)
                .filter(Report.student_id == payload.student_id, Report.release == payload.exam_id)
                .first()
            )
            if not report:
                report = Report(student_id=payload.student_id, release=payload.exam_id)
                db.add(report)
                db.flush()

            report.persona_template_id = persona_tmpl.id if persona_tmpl else None
            report.persona = persona_tmpl.student_label if persona_tmpl else payload.persona_code
            report.motivation_level = payload.motivation_level
            report.regulation_level = payload.regulation_level
            report.execution_level = payload.execution_level
            report.summary = payload.summary

            # Build new content map: indicator_id -> (analysis, suggestion)
            new_content: dict[int, tuple[str | None, str | None]] = {}
            for item in payload.strengths:
                ind_id = name_to_id.get(item.indicator_name, item.indicator_id)
                new_content[ind_id] = (item.analysis, None)
            for item in payload.weaknesses:
                ind_id = name_to_id.get(item.indicator_name, item.indicator_id)
                new_content[ind_id] = (item.analysis, item.suggestion)

            # Check if content unchanged
            current_rows = db.query(ReportIndicator).filter(
                ReportIndicator.report_id == report.id,
                ReportIndicator.is_current == True,
            ).all()
            if current_rows and len(current_rows) == len(new_content):
                current_map = {r.indicator_id: r for r in current_rows}
                if all(
                    current_map.get(ind_id) is not None
                    and current_map[ind_id].analysis == analysis
                    and current_map[ind_id].suggestion == suggestion
                    for ind_id, (analysis, suggestion) in new_content.items()
                ):
                    saved = [
                        SavedIndicatorAnalysis(
                            indicator_id=r.indicator_id,
                            indicator_name=next((n for n, i in name_to_id.items() if i == r.indicator_id), ""),
                            analysis=r.analysis,
                            suggestion=r.suggestion,
                            is_positive=r.is_positive,
                        )
                        for r in current_rows
                    ]
                    return ReportSaveResponse(
                        report_id=report.id,
                        student_id=payload.student_id,
                        exam_id=payload.exam_id,
                        indicators=saved,
                    )

            # Next version
            max_ver = db.query(func.max(ReportIndicator.version)).filter(
                ReportIndicator.report_id == report.id
            ).scalar() or 0
            next_version = max_ver + 1

            db.query(ReportIndicator).filter(
                ReportIndicator.report_id == report.id,
                ReportIndicator.is_current == True,
            ).update({"is_current": False}, synchronize_session=False)

            rows: list[ReportIndicator] = []
            for item in payload.strengths:
                ind_id = name_to_id.get(item.indicator_name, item.indicator_id)
                rows.append(ReportIndicator(
                    report_id=report.id,
                    indicator_id=ind_id,
                    analysis=item.analysis,
                    suggestion=None,
                    is_positive=True,
                    version=next_version,
                    is_current=True,
                ))
            for item in payload.weaknesses:
                ind_id = name_to_id.get(item.indicator_name, item.indicator_id)
                rows.append(ReportIndicator(
                    report_id=report.id,
                    indicator_id=ind_id,
                    analysis=item.analysis,
                    suggestion=item.suggestion,
                    is_positive=False,
                    version=next_version,
                    is_current=True,
                ))
            db.add_all(rows)

        saved = [
            SavedIndicatorAnalysis(
                indicator_id=r.indicator_id,
                indicator_name=next((n for n, i in name_to_id.items() if i == r.indicator_id), ""),
                analysis=r.analysis,
                suggestion=r.suggestion,
                is_positive=r.is_positive,
            )
            for r in rows
        ]
        return ReportSaveResponse(
            report_id=report.id,
            student_id=payload.student_id,
            exam_id=payload.exam_id,
            indicators=saved,
        )
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail={"message": "保存报告失败", "error": str(e)})


@router.get("/student/{student_id}", response_model=ReportGetResponse)
def get_report(
    student_id: int,
    exam_id: int,
    db: Session = Depends(get_db),
    current: Teacher = Depends(get_current_teacher),
):
    student = db.query(Student).filter(Student.id == student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    assert_student_class_access(current, student.class_id)

    report = (
        db.query(Report)
        .filter(Report.student_id == student_id, Report.release == exam_id)
        .first()
    )
    if not report:
        raise HTTPException(
            status_code=404,
            detail=f"未找到 student_id={student_id}, exam_id={exam_id} 的报告，请先生成报告",
        )

    rows = (
        db.query(ReportIndicator, Indicator)
        .join(Indicator, ReportIndicator.indicator_id == Indicator.id)
        .filter(ReportIndicator.report_id == report.id, ReportIndicator.is_current == True)
        .all()
    )
    indicators = [
        SavedIndicatorAnalysis(
            indicator_id=ind.id,
            indicator_name=ind.name,
            analysis=ri.analysis,
            suggestion=ri.suggestion,
            is_positive=ri.is_positive,
        )
        for ri, ind in rows
    ]

    persona = None
    if report.persona_template_id:
        tmpl = db.query(PersonaTemplate).filter(PersonaTemplate.id == report.persona_template_id).first()
        if tmpl:
            persona = PersonaResult(
                code=tmpl.code,
                teacher_label=tmpl.teacher_label,
                teacher_description=tmpl.teacher_description,
                student_label=tmpl.student_label,
                student_description=tmpl.student_description,
            )

    return ReportGetResponse(
        report_id=report.id,
        student_id=student_id,
        exam_id=exam_id,
        persona=persona,
        motivation_level=report.motivation_level,
        regulation_level=report.regulation_level,
        execution_level=report.execution_level,
        summary=report.summary,
        indicators=indicators,
    )


@router.get("/student/{student_id}/indicator-history", response_model=IndicatorHistoryResponse)
def get_indicator_history(
    student_id: int,
    exam_id: int,
    db: Session = Depends(get_db),
    current: Teacher = Depends(require_psych_or_above),
):
    student = db.query(Student).filter(Student.id == student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    assert_student_class_access(current, student.class_id)

    report = (
        db.query(Report)
        .filter(Report.student_id == student_id, Report.release == exam_id)
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    rows = (
        db.query(ReportIndicator, Indicator)
        .join(Indicator, ReportIndicator.indicator_id == Indicator.id)
        .filter(ReportIndicator.report_id == report.id)
        .order_by(ReportIndicator.indicator_id, ReportIndicator.version.desc())
        .all()
    )

    grouped: dict[int, dict] = {}
    for ri, ind in rows:
        if ind.id not in grouped:
            grouped[ind.id] = {
                "indicator_id": ind.id,
                "indicator_name": ind.name,
                "is_positive": ri.is_positive,
                "versions": [],
            }
        grouped[ind.id]["versions"].append(
            IndicatorVersion(
                version=ri.version,
                analysis=ri.analysis,
                suggestion=ri.suggestion,
                is_current=ri.is_current,
                created_at=ri.created_at,
            )
        )

    return IndicatorHistoryResponse(
        report_id=report.id,
        student_id=student_id,
        exam_id=exam_id,
        indicators=[IndicatorHistory(**v) for v in grouped.values()],
    )


def _save_student_report(db: Session, student_id: int, exam_id: int,
                          report_data: ReportGenerateResponse) -> None:
    """Persist a generated report to DB (extracted for reuse in background thread)."""
    persona_tmpl = db.query(PersonaTemplate).filter(
        PersonaTemplate.code == report_data.persona.code
    ).first()

    report = (
        db.query(Report)
        .filter(Report.student_id == student_id, Report.release == exam_id)
        .first()
    )
    if not report:
        report = Report(student_id=student_id, release=exam_id)
        db.add(report)
        db.flush()

    report.persona_template_id = persona_tmpl.id if persona_tmpl else None
    report.persona = persona_tmpl.student_label if persona_tmpl else report_data.persona.code
    report.motivation_level = report_data.persona.code[0]
    report.regulation_level = report_data.persona.code[1]
    report.execution_level = report_data.persona.code[2]
    report.summary = report_data.summary

    max_ver = db.query(func.max(ReportIndicator.version)).filter(
        ReportIndicator.report_id == report.id
    ).scalar() or 0
    next_version = max_ver + 1

    db.query(ReportIndicator).filter(
        ReportIndicator.report_id == report.id,
        ReportIndicator.is_current == True,
    ).update({"is_current": False}, synchronize_session=False)

    for ind in report_data.strengths:
        db.add(ReportIndicator(
            report_id=report.id,
            indicator_id=ind.indicator_id,
            analysis=ind.analysis,
            suggestion=None,
            is_positive=True,
            version=next_version,
            is_current=True,
        ))
    for ind in report_data.weaknesses:
        db.add(ReportIndicator(
            report_id=report.id,
            indicator_id=ind.indicator_id,
            analysis=ind.analysis,
            suggestion=ind.suggestion,
            is_positive=False,
            version=next_version,
            is_current=True,
        ))
    db.commit()


def _run_batch_job(job_id: int, exam_id: int, class_id: int | None,
                   student_ids: list[int] | None = None) -> None:
    """Background thread: generate & save reports for all students, update job progress."""
    from app.db.session import SessionLocal
    db = SessionLocal()
    try:
        # Load students
        query = db.query(Student)
        if student_ids:
            query = query.filter(Student.id.in_(student_ids))
        elif class_id:
            query = query.filter(Student.class_id == class_id)
        students = query.all()

        # Update job to running
        job = db.query(BatchJob).filter(BatchJob.id == job_id).first()
        job.status = "running"
        job.total = len(students)
        job.updated_at = datetime.utcnow()
        db.commit()

        errors = []
        for student in students:
            try:
                report_data = _generate_for_student(db, student.id, exam_id)
                _save_student_report(db, student.id, exam_id, report_data)
                job = db.query(BatchJob).filter(BatchJob.id == job_id).first()
                job.success += 1
            except Exception as e:
                errors.append({"student_id": student.id, "student_name": student.name,
                                "error": str(e)})
                job = db.query(BatchJob).filter(BatchJob.id == job_id).first()
                job.failed += 1

            job.errors = json.dumps(errors, ensure_ascii=False)
            job.updated_at = datetime.utcnow()
            db.commit()

        # Mark done
        job = db.query(BatchJob).filter(BatchJob.id == job_id).first()
        job.status = "done"
        job.updated_at = datetime.utcnow()
        db.commit()
    except Exception as e:
        try:
            job = db.query(BatchJob).filter(BatchJob.id == job_id).first()
            if job:
                job.status = "failed"
                job.errors = json.dumps([{"error": str(e)}], ensure_ascii=False)
                job.updated_at = datetime.utcnow()
                db.commit()
        except Exception:
            pass
    finally:
        db.close()


@router.get("/student-report-status")
def get_student_report_status(
    exam_id: int,
    class_id: int | None = None,
    db: Session = Depends(get_db),
    current: Teacher = Depends(require_psych_or_above),
):
    """Return all students with flags: has_report and data_changed.

    data_changed=True means scores were (re)computed after the report was generated,
    so the report may be stale and regeneration makes sense.
    data_changed=False means the underlying data is the same as when the report was built.
    """
    query = db.query(Student, Class).join(Class, Student.class_id == Class.id)
    if class_id:
        query = query.filter(Student.class_id == class_id)
    rows = query.order_by(Class.name, Student.name).all()

    # When scores were last computed for this exam
    exam = db.query(Exam).filter(Exam.id == exam_id).first()
    scores_computed_at = exam.scores_computed_at if exam else None

    # Fetch report creation times per student for this exam
    report_rows = (
        db.query(Report.student_id, Report.created_at)
        .filter(Report.release == exam_id)
        .all()
    )
    report_map: dict[int, object] = {r.student_id: r.created_at for r in report_rows}

    result = []
    for student, cls in rows:
        if current.role == "class_teacher" and current.class_id != student.class_id:
            continue
        has_report = student.id in report_map
        if not has_report:
            data_changed = True   # no report → generation is always meaningful
        elif scores_computed_at and report_map[student.id]:
            data_changed = scores_computed_at > report_map[student.id]
        else:
            data_changed = False  # can't determine → assume data unchanged
        result.append({
            "student_id": student.id,
            "student_name": student.name,
            "class_id": student.class_id,
            "class_name": cls.name,
            "has_report": has_report,
            "data_changed": data_changed,
        })
    return result


@router.post("/batch-generate")
def batch_generate_reports(
    payload: BatchGenerateRequest,
    db: Session = Depends(get_db),
    current: Teacher = Depends(require_psych_or_above),
):
    """Start a background batch-generate job and return job_id immediately."""
    query = db.query(Student)
    if payload.student_ids:
        query = query.filter(Student.id.in_(payload.student_ids))
    elif payload.class_id:
        query = query.filter(Student.class_id == payload.class_id)
    students = query.all()
    if not students:
        raise HTTPException(status_code=404, detail="未找到学生")

    student_id_list = [s.id for s in students]
    job = BatchJob(
        exam_id=payload.exam_id,
        class_id=payload.class_id,
        status="pending",
        total=len(students),
        student_ids=json.dumps(student_id_list),
    )
    db.add(job)
    db.commit()
    db.refresh(job)

    t = threading.Thread(
        target=_run_batch_job,
        args=(job.id, payload.exam_id, payload.class_id, payload.student_ids),
        daemon=True,
    )
    t.start()

    return {"job_id": job.id, "total": len(students), "status": "pending"}


@router.get("/batch-jobs/{job_id}")
def get_batch_job(
    job_id: int,
    db: Session = Depends(get_db),
    current: Teacher = Depends(require_psych_or_above),
):
    """Poll the status of a batch-generate job."""
    job = db.query(BatchJob).filter(BatchJob.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {
        "job_id": job.id,
        "status": job.status,           # pending / running / done / failed
        "total": job.total,
        "success": job.success,
        "failed": job.failed,
        "errors": json.loads(job.errors) if job.errors else [],
        "updated_at": job.updated_at.isoformat() if job.updated_at else None,
    }


@router.get("/batch-jobs")
def list_batch_jobs(
    exam_id: int,
    db: Session = Depends(get_db),
    current: Teacher = Depends(require_psych_or_above),
):
    """List all batch jobs for an exam (newest first)."""
    jobs = (
        db.query(BatchJob)
        .filter(BatchJob.exam_id == exam_id)
        .order_by(BatchJob.created_at.desc())
        .limit(10)
        .all()
    )
    return [
        {
            "job_id": j.id,
            "status": j.status,
            "total": j.total,
            "success": j.success,
            "failed": j.failed,
            "created_at": j.created_at.isoformat() if j.created_at else None,
        }
        for j in jobs
    ]


@router.get("/batch-jobs-all")
def list_all_batch_jobs(
    db: Session = Depends(get_db),
    current: Teacher = Depends(require_psych_or_above),
):
    """List all batch jobs across all exams (newest first, max 100)."""
    rows = (
        db.query(BatchJob, Exam.name.label("exam_name"))
        .join(Exam, BatchJob.exam_id == Exam.id)
        .order_by(BatchJob.created_at.desc())
        .limit(100)
        .all()
    )

    # Bulk-fetch all student names needed across all jobs
    all_student_ids: set[int] = set()
    for j, _ in rows:
        if j.student_ids:
            all_student_ids.update(json.loads(j.student_ids))
    student_name_map: dict[int, str] = {}
    if all_student_ids:
        for s in db.query(Student).filter(Student.id.in_(all_student_ids)).all():
            student_name_map[s.id] = s.name

    result = []
    for j, exam_name in rows:
        ids: list[int] = json.loads(j.student_ids) if j.student_ids else []
        students = [{"student_id": sid, "student_name": student_name_map.get(sid, f"学生{sid}")} for sid in ids]
        result.append({
            "job_id": j.id,
            "exam_id": j.exam_id,
            "exam_name": exam_name,
            "status": j.status,
            "total": j.total,
            "success": j.success,
            "failed": j.failed,
            "errors": json.loads(j.errors) if j.errors else [],
            "student_ids": ids,
            "students": students,
            "created_at": j.created_at.isoformat() if j.created_at else None,
            "updated_at": j.updated_at.isoformat() if j.updated_at else None,
        })
    return result


# ── DOCX Export ───────────────────────────────────────────────

_LEVEL_CN = {"H": "H", "M": "M", "L": "L"}
_SYSTEM_CN = {"motivation": "动力系统", "regulation": "调控系统", "execution": "执行系统"}


def _set_run_font(run, size_pt: float, bold: bool = False,
                  color: RGBColor | None = None) -> None:
    """Set Times New Roman (ASCII) + 宋体 (East Asian) on a run."""
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    # ASCII / Latin font
    run.font.name = "Times New Roman"
    # East Asian font via XML (python-docx has no direct API)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")


def _add_para_bottom_border(paragraph, color: str = "003399",
                             size: int = 6, space: int = 1) -> None:
    """Add a bottom border line to a paragraph (for title underline)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_docx(student_name: str, student_no: str, class_name: str, report: Report,
                persona_tmpl: PersonaTemplate | None, indicators: list) -> bytes:
    """Build a report .docx for one student and return bytes."""
    doc = Document()

    # ── Page setup: US Letter, 1.25" left/right, 1" top/bottom ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    def _para(text: str = "", bold: bool = False, size: float = 12,
              color: RGBColor | None = None, align=WD_ALIGN_PARAGRAPH.LEFT,
              indent: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(("　　" if indent else "") + text)
        _set_run_font(run, size, bold=bold, color=color)

    def _heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        _set_run_font(run, 14, bold=True, color=RGBColor(0x00, 0x70, 0xC0))

    def _label_value(label: str, value: str) -> None:
        p = doc.add_paragraph()
        r1 = p.add_run(label)
        _set_run_font(r1, 12, bold=True)
        r2 = p.add_run(value)
        _set_run_font(r2, 12)

    # ── Title ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("学生个性化分析报告")
    _set_run_font(title_run, 26, bold=True, color=RGBColor(0x00, 0x33, 0x99))
    _add_para_bottom_border(title_p, color="003399", size=6, space=1)

    # ── Student info ──
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(12)
    info_run = info_p.add_run(f"姓名：{student_name}    学号：{student_no}    班级：{class_name}")
    _set_run_font(info_run, 12, bold=True)

    doc.add_paragraph()  # blank line

    # ── 一、核心系统评估解释 ──
    _heading("一、核心系统评估解释")
    _para(
        "本次测评构建了三大核心系统：动力系统基于成长型思维与学习动机量表测定；"
        "调控系统涵盖注意控制能力与坚毅量表；执行系统则由学习策略量表表征。"
        "各指标评价均采用三级划分标准，即H（High）、M（Middle）、L（Low），"
        "分别对应高、中、低水平。",
        indent=True
    )

    # ── 二、典型画像分析 ──
    _heading("二、典型画像分析")

    mot = report.motivation_level or "M"
    reg = report.regulation_level or "M"
    exe = report.execution_level or "M"
    code_text = f"动力系统（{mot}）-调控系统（{reg}）-执行系统（{exe}）"
    _label_value("【画像特征码】: ", code_text)

    persona_label = persona_tmpl.student_label if persona_tmpl else (report.persona or "")
    _label_value("【典型画像类型】: ", persona_label)

    p_jd = doc.add_paragraph()
    r_jd = p_jd.add_run("【画像解读】:")
    _set_run_font(r_jd, 12, bold=True)

    if persona_tmpl and persona_tmpl.student_description:
        _para(persona_tmpl.student_description, indent=True)

    if report.summary:
        _para(report.summary, indent=True)

    # ── 三、优势项分析 ──
    _heading("三、具体指标优势项分析")
    strengths = [i for i in indicators if i.is_positive]
    for idx, ri in enumerate(strengths, 1):
        ind_name = ri.indicator_name if hasattr(ri, "indicator_name") else "—"
        score = ri.score_standardized if hasattr(ri, "score_standardized") else None
        level = ri.level if hasattr(ri, "level") else ""
        score_str = f"{score:.2f}" if score is not None else "—"
        p = doc.add_paragraph()
        r = p.add_run(f"{idx}. {ind_name} (得分: {score_str}) [{level}]")
        _set_run_font(r, 12, bold=True)
        if ri.analysis:
            _para(ri.analysis, indent=True)

    # ── 四、不足项分析 ──
    _heading("四、具体指标不足分析")
    weaknesses = [i for i in indicators if not i.is_positive]
    for idx, ri in enumerate(weaknesses, 1):
        ind_name = ri.indicator_name if hasattr(ri, "indicator_name") else "—"
        score = ri.score_standardized if hasattr(ri, "score_standardized") else None
        level = ri.level if hasattr(ri, "level") else ""
        score_str = f"{score:.2f}" if score is not None else "—"
        p = doc.add_paragraph()
        r = p.add_run(f"{idx}. {ind_name} (得分: {score_str}) [{level}]")
        _set_run_font(r, 12, bold=True)
        if ri.analysis:
            _para(ri.analysis, indent=True)

    # ── 五、改进建议 ──
    _heading("五、针对性改进建议")
    for idx, ri in enumerate(weaknesses, 1):
        ind_name = ri.indicator_name if hasattr(ri, "indicator_name") else "—"
        p = doc.add_paragraph()
        r = p.add_run(f"{idx}. 针对[{ind_name}]的建议:")
        _set_run_font(r, 12, bold=True)
        if ri.suggestion:
            _para(ri.suggestion, indent=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _get_full_report_data(db: Session, student_id: int, exam_id: int):
    """Fetch all data needed to build a docx for one student."""
    student = db.query(Student).filter(Student.id == student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")

    cls = db.query(Class).filter(Class.id == student.class_id).first()
    class_name = cls.name if cls else ""

    report = (
        db.query(Report)
        .filter(Report.student_id == student_id, Report.release == exam_id)
        .first()
    )
    if not report:
        raise HTTPException(status_code=404, detail="报告不存在，请先生成并保存报告")

    persona_tmpl = None
    if report.persona_template_id:
        persona_tmpl = db.query(PersonaTemplate).filter(
            PersonaTemplate.id == report.persona_template_id
        ).first()

    # Load saved indicators with score info
    rows = (
        db.query(ReportIndicator, Indicator)
        .join(Indicator, ReportIndicator.indicator_id == Indicator.id)
        .filter(ReportIndicator.report_id == report.id, ReportIndicator.is_current == True)
        .all()
    )

    class _RI:
        pass

    indicators = []
    for ri, ind in rows:
        obj = _RI()
        obj.indicator_name = ind.name
        obj.analysis = ri.analysis
        obj.suggestion = ri.suggestion
        obj.is_positive = ri.is_positive
        # Fetch score for this indicator
        score_row = db.query(ScoreStudent).filter(
            ScoreStudent.student_id == student_id,
            ScoreStudent.exam_id == exam_id,
            ScoreStudent.indicator_id == ind.id,
        ).first()
        obj.score_standardized = score_row.score_standardized if score_row else None
        obj.level = (
            "H" if (obj.score_standardized or 0) >= 0.67
            else "L" if (obj.score_standardized or 0) <= -0.67
            else "M"
        )
        indicators.append(obj)

    # Sort: strengths descending by score, weaknesses ascending
    strengths = sorted([i for i in indicators if i.is_positive],
                       key=lambda x: x.score_standardized or 0, reverse=True)
    weaknesses = sorted([i for i in indicators if not i.is_positive],
                        key=lambda x: x.score_standardized or 0)
    ordered = strengths + weaknesses

    return student, class_name, report, persona_tmpl, ordered


@router.get("/student/{student_id}/export-docx")
def export_single_docx(
    student_id: int,
    exam_id: int,
    db: Session = Depends(get_db),
    current: Teacher = Depends(get_current_teacher),
):
    """Export a single student's saved report as .docx."""
    student, class_name, report, persona_tmpl, indicators = _get_full_report_data(
        db, student_id, exam_id
    )
    assert_student_class_access(current, student.class_id)

    docx_bytes = _build_docx(
        student_name=student.name,
        student_no=str(student.id),
        class_name=class_name,
        report=report,
        persona_tmpl=persona_tmpl,
        indicators=indicators,
    )

    filename = f"{student.name}({student.id})-{class_name}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


class BatchExportRequest(BaseModel):
    exam_id: int
    student_ids: List[int]


@router.post("/batch-export-docx")
def batch_export_docx(
    payload: BatchExportRequest,
    db: Session = Depends(get_db),
    current: Teacher = Depends(require_psych_or_above),
):
    """Export selected students' saved reports as a ZIP of .docx files."""
    students = db.query(Student).filter(Student.id.in_(payload.student_ids)).all()
    if not students:
        raise HTTPException(status_code=404, detail="未找到学生")

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for student in students:
            try:
                _, class_name, report, persona_tmpl, indicators = _get_full_report_data(
                    db, student.id, payload.exam_id
                )
                docx_bytes = _build_docx(
                    student_name=student.name,
                    student_no=str(student.id),
                    class_name=class_name,
                    report=report,
                    persona_tmpl=persona_tmpl,
                    indicators=indicators,
                )
                filename = f"{student.name}({student.id})-{class_name}.docx"
                zf.writestr(filename, docx_bytes)
            except HTTPException:
                pass  # skip students without saved reports

    zip_buf.seek(0)
    return Response(
        content=zip_buf.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(f'reports_exam{payload.exam_id}.zip')}"},
    )
